from copy import deepcopy
from docx import Document
from romans import *
import docx2txt
import os
import re
import sys

path = "test.docx"   # sample file
# path = input("Write the path to a file: ")  # way with name or just name if file in project

# TODO: wrap everything in class


def fs(path_to_file):  # function witch split way, name and extension of file
    base = os.path.basename(path_to_file)
    way = os.path.split(path_to_file)[0]
    name = os.path.splitext(base)[0]
    ext = os.path.splitext(base)[1]
    return [way, name, ext]


def check_open(path_to_file):  # right file extension or not?
    p = path_to_file
    extension = fs(p)[2]
    if extension != '.docx':
        print("Wrong file, try again")
        sys.exit()
    else:
        print("Successful file opening")


def doc_text(path_to_file):  # extract text from the file
    p = path_to_file
    text = docx2txt.process(p)
    new_file = "{0}{1}_text.doc".format(fs(p)[0], fs(p)[1])
    text_file = open(new_file, 'w', encoding="utf8")
    text_file.write(text)
    text_file.close()
    print("Text saved")


def doc_tables(path_to_file):  # extract tables from the file
    p = path_to_file
    doc = Document(p)
    tabs = doc.tables
    if len(tabs) == 0:
        print("No Tables in the document")
    else:
        new_file = "{0}{1}_tables.docx".format(fs(p)[0], fs(p)[1])
        new_doc = Document()
        for t in range(len(tabs)):
            par = new_doc.add_paragraph(str("Таблица {}:".format(t+1)))
            tbl = deepcopy(tabs[t]._tbl)
            par._p.addnext(tbl)
        new_doc.save(new_file)
        print("Tables saved")


def replace_minus(lst):
    for i in lst:
        if len(i) > 1 and i[1:2].isdigit():
            lst[lst.index(i)] = re.sub('^-\S+$', 'минус {}'.format(i[1:]), i)
    return lst


def plus2text(path_to_file):
    p = path_to_file
    text_file = "{0}{1}_text.doc".format(fs(p)[0], fs(p)[1])
    check_lst = []
    with open(text_file, 'r', encoding="utf8") as file_read:
        all_lines = list(line for line in file_read)
    with open(text_file, 'w', encoding="utf8") as file_write:
        for line in all_lines:
            new_line = line.replace("+", "плюс ")
            check_lst.append(len(re.findall("плюс", new_line)))
            file_write.write(new_line)
    if len(check_lst) >= 1:
        print("Successful plus to text converting, count:", sum(check_lst))
    else:
        print("No + in the document")


def minus2text(path_to_file):
    p = path_to_file
    text_file = "{0}{1}_text.doc".format(fs(p)[0], fs(p)[1])
    check_lst = []
    with open(text_file, 'r', encoding="utf8") as file_read:
        all_lines = list(line for line in file_read)
    with open(text_file, 'w', encoding="utf8") as file_write:
        for line in all_lines:
            new_line = " ".join(replace_minus(line.split()))+"\n"
            check_lst.append(len(re.findall("минус", new_line)))
            file_write.write(new_line)
    if len(check_lst) >= 1:
        print("Successful minus to text converting, count:", sum(check_lst))
    else:
        print("No - in the document")


def del_pages(path_to_file):
    p = path_to_file
    text_file = "{0}{1}_text.doc".format(fs(p)[0], fs(p)[1])
    check_lst = []
    with open(text_file, 'r', encoding="utf8") as file_read:
        all_lines = list(line for line in file_read)
    with open(text_file, 'w', encoding="utf8") as file_write:
        for line in all_lines:
            if len(line.split()) == 1 and line.split()[0].isdigit():
                check_lst.append(True)
                continue
            if len(line.split()) > 1 and len(line.split()) < 11 and line.split()[1].isdigit():
                check_lst.append(True)
                continue
            else:
                file_write.write(line)
    if check_lst.count(True) >= 1:
        print("Successful delete page numbers, count:", check_lst.count(True))
    else:
        print("No Page Numbers in the document")


def convert_roman(path_to_file):
    p = path_to_file
    text_file = "{0}{1}_text.doc".format(fs(p)[0], fs(p)[1])
    count = 0
    with open(text_file, 'r', encoding="utf8") as file_read:
        all_lines = list(line for line in file_read)
    with open(text_file, 'w', encoding="utf8") as file_write:
        print("Process of converting roman numbers:")
        for line in all_lines:
            count = count + 1
            processing = "{} of {}".format(count, len(all_lines))
            print(processing)
            new_line = " ".join(rom2text2(rom2text(line.split())))+"\n"
            file_write.write(new_line)
        print("Successful Roman to text converting")

check_open(path)
doc_text(path)
doc_tables(path)
plus2text(path)
minus2text(path)
del_pages(path)
convert_roman(path)
